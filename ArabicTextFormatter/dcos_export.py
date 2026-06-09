# dcos_export.py - محرك التصدير (DOCX First)
import io
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import List, Dict, Any

def add_page_number_fields(run):
    """إضافة حقل رقم الصفحة في docx"""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    
    t = OxmlElement('w:t')
    t.text = "1"
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(t)
    run._r.append(fldChar3)

def export_to_docx(doc_model: List[Dict[str, Any]], formatting: Dict[str, Any]) -> bytes:
    """توليد ملف DOCX احترافي بالكامل انطلاقاً من الـ Document Model"""
    doc = Document()
    
    # 1. إعدادات تخطيط الصفحة وهوامشها
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    
    # الترويسة والتذييل
    start_second = formatting.get("start_from_second_page", True)
    if start_second:
        section.different_first_page_header_footer = True
        
    header_text = formatting.get("header_text", "")
    if header_text:
        header = section.header
        hp = header.paragraphs[0]
        hp.text = header_text
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    footer_text = formatting.get("footer_text", "")
    page_numbers = formatting.get("page_numbers", True)
    
    if footer_text or page_numbers:
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if footer_text:
            fp.text = footer_text + " - "
        if page_numbers:
            run = fp.add_run()
            add_page_number_fields(run)

    # 2. إعدادات الخطوط الافتراضية
    style = doc.styles['Normal']
    font = style.font
    font.name = formatting.get("font_family", "Traditional Arabic")
    font.size = Pt(formatting.get("font_size", 14))
    
    # دوال المحاذاة والتباعد
    def get_align(align_str):
        return {
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
        }.get(align_str, WD_ALIGN_PARAGRAPH.RIGHT)
        
    align_default = get_align(formatting.get("alignment", "right"))
    line_spacing_val = formatting.get("line_spacing", 1.5)
    space_after_val = Pt(formatting.get("space_after", 10))
    space_before_val = Pt(formatting.get("space_before", 0))
    
    # 3. بناء المستند عنصراً فعنصراً
    for item in doc_model:
        i_type = item["type"]
        
        if i_type == "empty_line":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            continue
            
        elif i_type == "heading":
            level = item.get("level", 1)
            # استخراج إعدادات التنسيق للعنوان المحدد
            h_config = formatting.get(f"heading{level}", {
                "size": 22 - (level * 2),
                "color": "#1e3c72" if level == 1 else "#333333",
                "alignment": "right",
                "weight": "bold",
                "underline": False
            })
            
            p = doc.add_paragraph()
            p.alignment = get_align(h_config.get("alignment", "right"))
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            
            run = p.add_run(item.get("text", ""))
            run.font.name = formatting.get("font_family", "Traditional Arabic")
            run.font.size = Pt(h_config.get("size", 16))
            run.font.color.rgb = RGBColor.from_string(h_config.get("color", "#1e3c72").replace("#", ""))
            run.font.underline = h_config.get("underline", False)
            
            if h_config.get("weight", "bold") in ["bold", "semi-bold"]:
                run.font.bold = True
                
        elif i_type == "bullet_list_item":
            p = doc.add_paragraph()
            p.alignment = align_default
            p.paragraph_format.line_spacing = line_spacing_val
            p.paragraph_format.space_after = space_after_val
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            
            bullet_sym = formatting.get("bullet_symbol", "•")
            run = p.add_run(f"{bullet_sym} {item.get('text', '')}")
            run.font.name = formatting.get("font_family", "Traditional Arabic")
            run.font.size = Pt(formatting.get("font_size", 14))
            run.font.color.rgb = RGBColor.from_string(formatting.get("font_color", "#1a1a1a").replace("#", ""))
            
        elif i_type == "numbered_list_item":
            p = doc.add_paragraph()
            p.alignment = align_default
            p.paragraph_format.line_spacing = line_spacing_val
            p.paragraph_format.space_after = space_after_val
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            
            prefix = item.get("prefix", "1.")
            run = p.add_run(f"{prefix} {item.get('text', '')}")
            run.font.name = formatting.get("font_family", "Traditional Arabic")
            run.font.size = Pt(formatting.get("font_size", 14))
            run.font.color.rgb = RGBColor.from_string(formatting.get("font_color", "#1a1a1a").replace("#", ""))
            
        elif i_type == "quote":
            p = doc.add_paragraph()
            p.alignment = get_align("center")
            p.paragraph_format.line_spacing = line_spacing_val
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.right_indent = Inches(0.5)
            
            run = p.add_run(item.get("text", ""))
            run.font.name = formatting.get("font_family", "Traditional Arabic")
            run.font.size = Pt(formatting.get("font_size", 12))
            run.font.italic = True
            run.font.color.rgb = RGBColor.from_string("555555")
            
        else: # paragraph
            p = doc.add_paragraph()
            p.alignment = align_default
            p.paragraph_format.line_spacing = line_spacing_val
            p.paragraph_format.space_before = space_before_val
            p.paragraph_format.space_after = space_after_val
            
            run = p.add_run(item.get("text", ""))
            run.font.name = formatting.get("font_family", "Traditional Arabic")
            run.font.size = Pt(formatting.get("font_size", 14))
            run.font.color.rgb = RGBColor.from_string(formatting.get("font_color", "#1a1a1a").replace("#", ""))
            
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()

def export_to_markdown(doc_model: List[Dict[str, Any]]) -> str:
    """تصدير كـ Markdown"""
    lines = []
    for item in doc_model:
        i_type = item["type"]
        if i_type == "empty_line":
            lines.append("")
        elif i_type == "heading":
            level = item.get("level", 1)
            lines.append(f"{'#' * level} {item.get('text', '')}")
        elif i_type == "bullet_list_item":
            lines.append(f"* {item.get('text', '')}")
        elif i_type == "numbered_list_item":
            prefix = item.get("prefix", "1.")
            lines.append(f"{prefix} {item.get('text', '')}")
        elif i_type == "quote":
            lines.append(f"> {item.get('text', '')}")
        else:
            lines.append(item.get("text", ""))
    return "\n".join(lines)

def export_to_html(doc_model: List[Dict[str, Any]], formatting: Dict[str, Any]) -> str:
    """توليد كود HTML متكامل للمعاينة والطباعة المباشرة"""
    font_fam = formatting.get("font_family", "Cairo")
    font_sz = formatting.get("font_size", 14)
    font_col = formatting.get("font_color", "#1a1a1a")
    alignment = formatting.get("alignment", "right")
    line_sp = formatting.get("line_spacing", 1.5)
    space_after = formatting.get("space_after", 10)
    
    html_out = [
        "<!DOCTYPE html>",
        "<html lang='ar' dir='rtl'>",
        "<head>",
        "<meta charset='UTF-8'>",
        f"<style>",
        f"  body {{ font-family: '{font_fam}', sans-serif; font-size: {font_sz}pt; color: {font_col}; line-height: {line_sp}; text-align: {alignment}; direction: rtl; padding: 20px; }}",
        f"  .heading {{ margin-top: 20px; margin-bottom: 8px; font-weight: bold; }}",
        f"  .bullet-item {{ margin-right: 20px; margin-bottom: {space_after}px; }}",
        f"  .number-item {{ margin-right: 20px; margin-bottom: {space_after}px; }}",
        f"  .paragraph {{ margin-bottom: {space_after}px; }}",
        f"  .quote {{ font-style: italic; color: #555; margin: 15px 40px; border-right: 3px solid #ccc; padding-right: 15px; text-align: center; }}",
        f"</style>",
        "</head>",
        "<body>"
    ]
    
    weight_map = { 'bold': '700', 'semi-bold': '600', 'regular': '400' }
    
    for item in doc_model:
        i_type = item["type"]
        if i_type == "empty_line":
            html_out.append("<div style='height: 10px;'></div>")
        elif i_type == "heading":
            level = item.get("level", 1)
            h_config = formatting.get(f"heading{level}", {
                "size": 22 - (level * 2),
                "color": "#1e3c72" if level == 1 else "#333333",
                "alignment": "right",
                "weight": "bold",
                "underline": False
            })
            fw = weight_map.get(h_config.get("weight", "bold"), "bold")
            td = "underline" if h_config.get("underline", False) else "none"
            html_out.append(f"<h{level} class='heading' style='font-size:{h_config.get('size', 16)}pt; color:{h_config.get('color', '#1e3c72')}; text-align:{h_config.get('alignment', 'right')}; font-weight:{fw}; text-decoration:{td};'>{item.get('text', '')}</h{level}>")
        elif i_type == "bullet_list_item":
            bullet_sym = formatting.get("bullet_symbol", "•")
            html_out.append(f"<div class='bullet-item'>{bullet_sym} {item.get('text', '')}</div>")
        elif i_type == "numbered_list_item":
            prefix = item.get("prefix", "1.")
            html_out.append(f"<div class='number-item'>{prefix} {item.get('text', '')}</div>")
        elif i_type == "quote":
            html_out.append(f"<div class='quote'>{item.get('text', '')}</div>")
        else:
            html_out.append(f"<div class='paragraph'>{item.get('text', '')}</div>")
            
    html_out.append("</body></html>")
    return "\n".join(html_out)
